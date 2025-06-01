import streamlit as st
st.set_page_config(page_title="AutoMLify Dashboard", layout="wide")
import pandas as pd
from io import BytesIO
import docx
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from preprocessing.preprocess_pipeline import handle_missing_values, balance_dataset, encode_and_scale_features
from eda.eda_pipeline import run_eda
from feature_engineering.feature_pipeline import (
    remove_low_variance_features,
    remove_highly_correlated_features,
    select_important_features, scale_features
)
from sklearn.preprocessing import LabelEncoder

from model_training.training_pipeline import auto_mode, manual_mode, detect_task_type
from evaluation.evaluation_pipeline import (
    evaluate_model,
    generate_report_dict
)


st.title("🧠 AutoMLify Dashboard")

st.sidebar.markdown("## 📂 Navigation")
page = st.sidebar.radio("Go to", [
    "Upload Dataset", "Preprocessing", "EDA",
    "Feature Engineering", "Model Training and Evaluation", "Evaluation"])

if page == "Upload Dataset":
    st.header("📤 Upload Your Dataset")
    uploaded_file = st.file_uploader("Upload CSV file", type=["csv"])
    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        st.session_state.df = df
        st.success("✅ Dataset uploaded successfully!")
        with st.expander("🔍 View Raw Data"):
            st.dataframe(df.head())

elif page == "Preprocessing":
    if "df" in st.session_state:
        df = st.session_state.df.copy()
        df_cleaned = handle_missing_values(df)
        st.session_state.df_cleaned = df_cleaned

        st.write("✅ Cleaned Dataset:")
        st.dataframe(df_cleaned.head())

        target_col = st.selectbox("🎯 Select Target Column (for balancing)", df_cleaned.columns)
        if target_col:
            X = df_cleaned.drop(columns=[target_col])
            y = df_cleaned[target_col]

            if y.dtype == 'object':
                label_encoder = LabelEncoder()
                y = label_encoder.fit_transform(y)

            X_balanced, y_balanced = balance_dataset(X, y)
            X_encoded_scaled = encode_and_scale_features(X_balanced)

            st.session_state.X_processed = X_encoded_scaled
            st.session_state.y_processed = y_balanced

            st.success("✅ Preprocessing complete. Proceed to EDA or Modeling.")
    else:
        st.warning("⚠️ Please upload a dataset first.")

elif page == "EDA":
    if "df_cleaned" in st.session_state:
        st.subheader("🔎 Exploratory Data Analysis")
        df = st.session_state.df_cleaned
        run_eda(df)
    else:
        st.warning("⚠️ Please run preprocessing first.")

elif page == "Feature Engineering":
    if "X_processed" in st.session_state and "y_processed" in st.session_state:
        st.subheader("🧪 Feature Engineering")
        X = st.session_state.X_processed.copy()
        y = st.session_state.y_processed

        X = remove_low_variance_features(X)
        X = remove_highly_correlated_features(X)

        if X.shape[0] > 0:
            method = st.selectbox("🎯 Feature selection method", ["f_classif"])
            k = st.slider("🔢 Number of top features to select", 1, min(20, X.shape[1]), 5)
            X = select_important_features(X, y, method, k)
        else:
            st.warning("⚠️ No data samples available for feature selection. Please check your preprocessing.")

        X_scaled = scale_features(X)

        st.session_state.X_final = X_scaled
        st.session_state.y_final = y

        st.success("✅ Feature engineering complete. Ready for model training!")
    else:
        st.warning("⚠️ Please complete preprocessing first.")

elif page == "Model Training and Evaluation":
    if "X_final" in st.session_state and "y_final" in st.session_state:
        X = st.session_state.X_final
        y = st.session_state.y_final

        task_type = detect_task_type(y)
        st.subheader("⚙️ Model Training")
        mode = st.radio("Choose Training Mode", ["Auto", "Manual"])

        with st.spinner("Training your model..."):
            if mode == "Auto":
                auto_mode(X, y, task_type)
            else:
                manual_mode(X, y, task_type)
        st.success("✅ Model training completed.")
    else:
        st.warning("⚠️ Please complete feature engineering before training.")

elif page == "Evaluation":
    if all(k in st.session_state for k in ["trained_model", "X_test", "y_test", "task_type"]):
        st.subheader("📊 Model Evaluation")
        model = st.session_state.trained_model
        X_test = st.session_state.X_test
        y_test = st.session_state.y_test
        task_type = st.session_state.task_type

        metrics_dict = evaluate_model(model, X_test, y_test, task_type)

        report = generate_report_dict(
            model_name=type(model).__name__,
            model_params=model.get_params(),
            metrics_dict=metrics_dict,
            task_type=task_type
        )

        for section, content in report.items():
            st.markdown(f"### {section}")
            st.write(content)

        report_df = pd.DataFrame(list(report.items()), columns=["Section", "Details"])
        csv = report_df.to_csv(index=False).encode('utf-8')
        st.download_button("📥 Download CSV", data=csv, file_name='evaluation_report.csv', mime='text/csv')

        # Word
        word_buffer = BytesIO()
        doc = docx.Document()
        doc.add_heading("Evaluation Report", 0)
        for section, detail in report.items():
            doc.add_heading(section, level=1)
            doc.add_paragraph(str(detail))
        doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button("📄 Download Word", data=word_buffer, file_name="evaluation_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # PDF
        pdf_buffer = BytesIO()
        doc_pdf = SimpleDocTemplate(pdf_buffer)
        styles = getSampleStyleSheet()
        story = [Paragraph("Evaluation Report", styles["Title"])]
        for section, detail in report.items():
            story.append(Paragraph(f"<b>{section}</b>", styles["Heading2"]))
            story.append(Paragraph(str(detail), styles["BodyText"]))
        doc_pdf.build(story)
        pdf_buffer.seek(0)
        st.download_button("📝 Download PDF", data=pdf_buffer, file_name="evaluation_report.pdf", mime="application/pdf")
    else:
        st.warning("⚠️ Please train a model first before accessing evaluation.")
